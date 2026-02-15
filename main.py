import base64
import json
import os
import re
import uuid
from io import BytesIO

import pdfplumber
import google.generativeai as genai
from docx import Document as DocxDocument
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# ─────────────────────────────────────────────
# ENV SETUP
# ─────────────────────────────────────────────
# load_dotenv()
# # Make sure the environment variable is loaded correctly
# # GOOGLE_API_KEY= st.secrets["GOOGLE_API_KEY"]
load_dotenv()

api_key = os.getenv("GOOGLE_API_KEY")
if not api_key:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")
genai.configure(api_key=api_key)

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
SIMILARITY_THRESHOLD = 35    # display score 0-100; refuse if best chunk BELOW this
LOW_CONFIDENCE_CUTOFF = 30   # refuse answer if LLM confidence below this

# In-memory session store — keyed by session_id (uuid)
# Each session: { "vector_store": FAISS, "file_names": [...], "chunk_count": int }
SESSION_STORE: dict = {}

# ─────────────────────────────────────────────
# APP
# ─────────────────────────────────────────────
app = FastAPI(
    title="TMS Doc Intelligence API",
    description="RAG-powered logistics document Q&A and extraction API",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─────────────────────────────────────────────
# REQUEST / RESPONSE MODELS
# ─────────────────────────────────────────────
class UploadRequest(BaseModel):
    """
    Accept one or more files as base64 strings.
    file_names is optional but helpful for logging.
    """
    files_base64: list[str]
    file_names: list[str] = []

class UploadResponse(BaseModel):
    session_id: str
    message: str
    chunk_count: int
    file_names: list[str]

class AskRequest(BaseModel):
    session_id: str
    question: str

class SourceChunk(BaseModel):
    chunk_index: int
    text: str
    similarity_score: int  # 0-100

class AskResponse(BaseModel):
    session_id: str
    question: str
    answer: str
    confidence: int | None
    sources: list[SourceChunk]
    guardrail_triggered: bool
    guardrail_reason: str | None

class ExtractRequest(BaseModel):
    session_id: str

class ExtractResponse(BaseModel):
    session_id: str
    data: dict

# ─────────────────────────────────────────────
# TEXT EXTRACTION HELPERS
# ─────────────────────────────────────────────
def detect_file_type(file_bytes: bytes, file_name: str = "") -> str:
    # First try magic bytes — most reliable
    if file_bytes[:4] == b"%PDF":
        return "pdf"
    elif file_bytes[:4] == b"PK\x03\x04":
        return "docx"
    
    # Fallback to file extension if bytes are ambiguous (e.g. TXT files have no magic bytes)
    if file_name:
        ext = file_name.rsplit(".", 1)[-1].lower()
        if ext in ("pdf",):
            return "pdf"
        elif ext in ("docx", "doc"):
            return "docx"
        elif ext in ("txt",):
            return "txt"
    
    # Last resort — try decoding as UTF-8 text
    try:
        file_bytes.decode("utf-8")
        return "txt"
    except UnicodeDecodeError:
        return "unknown"

def extract_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True)
            if page_text:
                text += page_text + "\n"
    return text

def extract_docx(file_bytes: bytes) -> str:
    text = ""
    doc = DocxDocument(BytesIO(file_bytes))
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"
    for table in doc.tables:
        text += "\n"
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text += " | ".join(row_text) + "\n"
    return text

def extract_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8").strip()

def extract_text_from_base64(file_base64: str, file_name: str = "") -> str:
    file_bytes = base64.b64decode(file_base64)
    file_type = detect_file_type(file_bytes, file_name)   # <-- pass name here
    if file_type == "pdf":
        return extract_pdf(file_bytes)
    elif file_type == "docx":
        return extract_docx(file_bytes)
    elif file_type == "txt":
        return extract_txt(file_bytes)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, TXT are supported.")

# ─────────────────────────────────────────────
# CHUNKING & VECTOR STORE
# ─────────────────────────────────────────────
def build_vector_store(text: str) -> tuple[FAISS, int]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunks = splitter.split_text(text)
    embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    return vector_store, len(chunks)

# ─────────────────────────────────────────────
# RETRIEVAL WITH SIMILARITY GUARDRAIL
# ─────────────────────────────────────────────
def retrieve_with_scores(question: str, vector_store: FAISS, k: int = 4):
    """
    Returns (docs, display_scores, passed_threshold, best_display_score).

    FAISS returns L2 distances (lower = better match).
    We convert to display score 0-100 (higher = better):
        display = max(0, round((1 - l2 / 2) * 100))

    Guardrail: refuse if best display score < SIMILARITY_THRESHOLD (35).
    """
    results = vector_store.similarity_search_with_score(question, k=k)
    docs = [r[0] for r in results]
    l2_scores = [r[1] for r in results]

    display_scores = [max(0, round((1 - s / 2) * 100)) for s in l2_scores]
    best_display_score = max(display_scores) if display_scores else 0
    passed = best_display_score >= SIMILARITY_THRESHOLD

    return docs, display_scores, passed, best_display_score

# ─────────────────────────────────────────────
# LLM CHAIN
# ─────────────────────────────────────────────
def get_qa_chain():
    prompt_template = """
You are a logistics document assistant. Answer ONLY from the provided context.

Rules:
- If the answer is clearly present: answer precisely.
- If the answer is partially present: answer what you can and note what is missing.
- If the answer is NOT in the context at all: respond exactly with "Not found in document."
- Do NOT generalize or use outside knowledge.
- Always end your response with exactly this line:
  Confidence: <number between 0 and 100>

Context:
{context}

Question:
{question}

Answer:
"""
    model = ChatGoogleGenerativeAI(model="models/gemma-3-27b-it", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

def get_extraction_chain():
    prompt_template = """
You are an expert at extracting structured data from logistics documents.
Extract the fields below strictly from the context. Return ONLY valid JSON, no explanation, no markdown fences.
If a field is not found return null for that field.

{{
  "shipment_id": null,
  "shipper": null,
  "consignee": null,
  "pickup_datetime": null,
  "delivery_datetime": null,
  "equipment_type": null,
  "mode": null,
  "rate": null,
  "currency": null,
  "weight": null,
  "carrier_name": null
}}

Context:
{context}

Question:
{question}

Answer (only JSON):
"""
    model = ChatGoogleGenerativeAI(model="models/gemma-3-27b-it", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

# ─────────────────────────────────────────────
# CONFIDENCE PARSING
# ─────────────────────────────────────────────
def parse_confidence(text: str) -> int | None:
    patterns = [
        r"confidence[^\d]{0,15}(\d{1,3})",
        r"(\d{1,3})\s*(?:%|percent)\s*confident",
        r"\[(\d{1,3})\]",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = int(m.group(1))
            if 0 <= val <= 100:
                return val
    return None

def clean_answer(raw: str) -> str:
    return re.sub(
        r"(confidence\s*(?:score)?[:\s\-]*\d{1,3}\s*%?\.?)",
        "",
        raw,
        flags=re.IGNORECASE
    ).strip()

# ─────────────────────────────────────────────
# ENDPOINTS
# ─────────────────────────────────────────────

@app.get("/")
def root():
    return {
        "status": "TMS Doc Intelligence API running",
        "endpoints": {
            "POST /upload": "Upload and process documents",
            "POST /ask":    "Ask a question about uploaded documents",
            "POST /extract":"Extract structured shipment data"
        }
    }


# ── POST /upload ──────────────────────────────
@app.post("/upload", response_model=UploadResponse)
def upload(request: UploadRequest):
    """
    Upload one or more documents as base64 strings.
    Returns a session_id to use in /ask and /extract.

    Body:
    {
        "files_base64": ["<base64_string>", ...],
        "file_names":   ["invoice.pdf", ...]   // optional
    }
    """
    if not request.files_base64:
        raise HTTPException(status_code=400, detail="No files provided.")

    combined_text = ""
    for i, file_b64 in enumerate(request.files_base64):
        try:
            # Pass the matching filename if provided, else empty string
            name = request.file_names[i] if i < len(request.file_names) else ""
            combined_text += extract_text_from_base64(file_b64, name) + "\n"  # <-- pass name
        except ValueError as e:
            raise HTTPException(status_code=400, detail=f"File {i+1}: {str(e)}")

    if not combined_text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from the uploaded files.")

    try:
        vector_store, chunk_count = build_vector_store(combined_text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to build vector index: {str(e)}")

    session_id = str(uuid.uuid4())
    SESSION_STORE[session_id] = {
        "vector_store": vector_store,
        "file_names": request.file_names or [f"file_{i+1}" for i in range(len(request.files_base64))],
        "chunk_count": chunk_count,
    }

    return UploadResponse(
        session_id=session_id,
        message="Documents processed and indexed successfully.",
        chunk_count=chunk_count,
        file_names=SESSION_STORE[session_id]["file_names"],
    )


# ── POST /ask ─────────────────────────────────
@app.post("/ask", response_model=AskResponse)
def ask(request: AskRequest):
    """
    Ask a natural language question about the uploaded document.
    Returns answer, confidence score, source chunks, and guardrail info.

    Body:
    {
        "session_id": "<uuid from /upload>",
        "question":   "What is the carrier rate?"
    }
    """
    if request.session_id not in SESSION_STORE:
        raise HTTPException(status_code=404, detail="session_id not found. Please call /upload first.")

    vector_store = SESSION_STORE[request.session_id]["vector_store"]

    # ── Retrieve with similarity scores ──
    docs, display_scores, passed_threshold, best_score = retrieve_with_scores(
        request.question, vector_store
    )

    source_chunks = [
        SourceChunk(
            chunk_index=i + 1,
            text=doc.page_content,
            similarity_score=score
        )
        for i, (doc, score) in enumerate(zip(docs, display_scores))
    ]

    # ── GUARDRAIL 1: Retrieval similarity too low ──
    if not passed_threshold:
        return AskResponse(
            session_id=request.session_id,
            question=request.question,
            answer="Answer refused: retrieval similarity too low. This question likely refers to content not in the document.",
            confidence=None,
            sources=source_chunks,
            guardrail_triggered=True,
            guardrail_reason=f"Best chunk similarity score {best_score}% is below threshold {SIMILARITY_THRESHOLD}%."
        )

    # ── Call LLM ──
    chain = get_qa_chain()
    response = chain(
        {"input_documents": docs, "question": request.question},
        return_only_outputs=True
    )
    raw_answer = response["output_text"]

    # ── GUARDRAIL 2: Not found in document ──
    if "not found in document" in raw_answer.lower():
        return AskResponse(
            session_id=request.session_id,
            question=request.question,
            answer="Not found in document.",
            confidence=0,
            sources=source_chunks,
            guardrail_triggered=True,
            guardrail_reason="LLM determined the answer is not present in the document context."
        )

    confidence = parse_confidence(raw_answer)
    answer = clean_answer(raw_answer)

    # ── GUARDRAIL 3: LLM confidence too low ──
    if confidence is not None and confidence < LOW_CONFIDENCE_CUTOFF:
        return AskResponse(
            session_id=request.session_id,
            question=request.question,
            answer=f"Answer withheld: confidence too low ({confidence}%). Try rephrasing your question.",
            confidence=confidence,
            sources=source_chunks,
            guardrail_triggered=True,
            guardrail_reason=f"LLM confidence {confidence}% is below minimum threshold {LOW_CONFIDENCE_CUTOFF}%."
        )

    return AskResponse(
        session_id=request.session_id,
        question=request.question,
        answer=answer,
        confidence=confidence,
        sources=source_chunks,
        guardrail_triggered=False,
        guardrail_reason=None
    )


# ── POST /extract ─────────────────────────────
@app.post("/extract", response_model=ExtractResponse)
def extract(request: ExtractRequest):
    """
    Extract structured shipment data from the uploaded document.
    Returns JSON with all 11 standard TMS fields. Missing fields are null.

    Body:
    {
        "session_id": "<uuid from /upload>"
    }
    """
    if request.session_id not in SESSION_STORE:
        raise HTTPException(status_code=404, detail="session_id not found. Please call /upload first.")

    vector_store = SESSION_STORE[request.session_id]["vector_store"]

    extraction_question = """
Extract these fields from the document. Return only valid JSON, no explanation.
Fields: shipment_id, shipper, consignee, pickup_datetime, delivery_datetime,
equipment_type, mode, rate, currency, weight, carrier_name.
"""

    # Fetch broad context for extraction
    docs_primary, _, _, _ = retrieve_with_scores(extraction_question, vector_store, k=4)
    docs_secondary, _, _, _ = retrieve_with_scores(
        "shipment shipper consignee carrier rate weight pickup delivery", vector_store, k=4
    )

    # Deduplicate by content
    seen = set()
    all_docs = []
    for doc in docs_primary + docs_secondary:
        if doc.page_content not in seen:
            seen.add(doc.page_content)
            all_docs.append(doc)

    chain = get_extraction_chain()
    response = chain(
        {"input_documents": all_docs, "question": extraction_question},
        return_only_outputs=True
    )
    raw = response["output_text"].strip()

    # Strip markdown fences if model ignores instructions
    raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"```$", "", raw).strip()

    # Find the JSON object
    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if json_match:
        raw = json_match.group(0)

    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        # Return nulls for all fields rather than failing hard
        parsed = {
            "shipment_id": None, "shipper": None, "consignee": None,
            "pickup_datetime": None, "delivery_datetime": None,
            "equipment_type": None, "mode": None, "rate": None,
            "currency": None, "weight": None, "carrier_name": None
        }

    return ExtractResponse(
        session_id=request.session_id,
        data=parsed
    )


# ─────────────────────────────────────────────
# RUN
# ─────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("fastmain:app", host="0.0.0.0", port=8000, reload=True)
