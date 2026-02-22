from docx import Document as DocxDocument
import streamlit as st
import os
import json
import re
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv
import google.generativeai as genai
from dotenv import load_dotenv
import google.generativeai as genai
import pdfplumber
GOOGLE_API_KEY= st.secrets["GOOGLE_API_KEY"]
# load_dotenv()

# ─────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────
SIMILARITY_THRESHOLD = 1.2   # FAISS L2 distance threshold; tune this based on your embeddings
LOW_CONFIDENCE_CUTOFF = 30

# api_key = os.getenv("GOOGLE_API_KEY")
# if api_key is None:
#     raise ValueError("GOOGLE_API_KEY environment variable is not set")
# genai.configure(api_key=api_key)

# ─────────────────────────────────────────────
# PAGE CONFIG & CUSTOM CSS
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="TMS Doc Intelligence",
    page_icon="🚚",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap');

    :root {
        --bg-base:        #080c14;
        --bg-surface:     #0d1220;
        --bg-elevated:    #111827;
        --bg-card:        #141d2e;
        --border:         #1e2d45;
        --border-bright:  #243650;
        --accent:         #00d4ff;
        --accent-dim:     #00a3c4;
        --accent-glow:    rgba(0, 212, 255, 0.12);
        --accent-glow2:   rgba(0, 212, 255, 0.06);
        --success:        #00e5a0;
        --success-bg:     rgba(0, 229, 160, 0.08);
        --warn:           #ffb347;
        --warn-bg:        rgba(255, 179, 71, 0.08);
        --danger:         #ff5c5c;
        --danger-bg:      rgba(255, 92, 92, 0.08);
        --text-primary:   #e8edf5;
        --text-secondary: #7a8fa8;
        --text-muted:     #3d5068;
        --font-main:      'Space Grotesk', sans-serif;
        --font-mono:      'JetBrains Mono', monospace;
        --radius-sm:      6px;
        --radius-md:      10px;
        --radius-lg:      14px;
    }

    html, body, [class*="css"] {
        font-family: var(--font-main);
    }

    /* ── Base ── */
    .stApp {
        background: var(--bg-base);
    }
    .main .block-container {
        padding: 2rem 2.5rem 3rem;
        max-width: 960px;
    }

    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {
        background: var(--bg-surface);
        border-right: 1px solid var(--border);
    }
    section[data-testid="stSidebar"] > div {
        padding: 1.5rem 1.2rem;
    }
    section[data-testid="stSidebar"] * {
        color: var(--text-primary) !important;
    }

    /* Sidebar title */
    .sidebar-brand {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 0.6rem 0 1.2rem;
    }
    .sidebar-brand-icon {
        font-size: 1.4rem;
        line-height: 1;
    }
    .sidebar-brand-text {
        font-family: var(--font-mono);
        font-size: 0.82rem;
        font-weight: 600;
        color: var(--accent) !important;
        letter-spacing: 0.5px;
        line-height: 1.3;
    }
    .sidebar-brand-sub {
        font-size: 0.65rem;
        color: var(--text-muted) !important;
        font-weight: 400;
        letter-spacing: 0.5px;
    }

    /* ── Top toolbar / deploy bar ── */
    header[data-testid="stHeader"] {
        background: var(--bg-surface) !important;
        border-bottom: 1px solid var(--border) !important;
    }
    header[data-testid="stHeader"] * {
        color: var(--text-secondary) !important;
    }
    /* The top-right deploy/menu toolbar */
    [data-testid="stToolbar"] {
        background: var(--bg-surface) !important;
        border-bottom: 1px solid var(--border) !important;
    }
    [data-testid="stToolbar"] button,
    [data-testid="stToolbar"] a {
        color: var(--text-secondary) !important;
    }
    /* Full-width white bar at very top */
    .stDeployButton, [data-testid="stDecoration"] {
        background: var(--bg-surface) !important;
        display: none !important;
    }
    /* App top chrome */
    #root > div:first-child > div:first-child {
        background: var(--bg-surface) !important;
    }
    /* Iframe / top bar */
    .viewerBadge_container__1QSob,
    .viewerBadge_link__1S137 {
        display: none !important;
    }

    /* ── Upload area ── */
    [data-testid="stFileUploader"] {
        border: 1.5px dashed var(--border-bright) !important;
        border-radius: var(--radius-md) !important;
        background: var(--bg-elevated) !important;
        transition: border-color 0.2s;
    }
    [data-testid="stFileUploader"]:hover {
        border-color: var(--accent-dim) !important;
    }
    /* Upload drop zone inner area */
    [data-testid="stFileUploaderDropzone"] {
        background: var(--bg-elevated) !important;
        border: 1.5px dashed var(--border-bright) !important;
        border-radius: var(--radius-md) !important;
    }
    [data-testid="stFileUploaderDropzone"]:hover {
        border-color: var(--accent-dim) !important;
        background: var(--bg-card) !important;
    }
    /* Upload zone text */
    [data-testid="stFileUploaderDropzoneInstructions"] {
        color: var(--text-secondary) !important;
    }
    [data-testid="stFileUploaderDropzoneInstructions"] small,
    [data-testid="stFileUploaderDropzoneInstructions"] span {
        color: var(--text-muted) !important;
    }
    /* Browse files button inside uploader */
    [data-testid="stFileUploaderDropzone"] button {
        background: var(--bg-card) !important;
        color: var(--accent) !important;
        border: 1px solid var(--border-bright) !important;
        border-radius: var(--radius-sm) !important;
        font-family: var(--font-mono) !important;
        font-size: 0.75rem !important;
        font-weight: 500 !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover {
        background: var(--accent-glow) !important;
        border-color: var(--accent-dim) !important;
    }
    /* Uploaded file item */
    [data-testid="stFileUploaderFile"] {
        background: var(--bg-card) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-sm) !important;
    }
    [data-testid="stFileUploaderFile"] * {
        color: var(--text-secondary) !important;
    }
    /* Label above uploader */
    [data-testid="stFileUploader"] label {
        color: var(--text-secondary) !important;
        font-family: var(--font-mono) !important;
        font-size: 0.78rem !important;
    }

    /* Sidebar stats box */
    .stats-box {
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-radius: var(--radius-md);
        padding: 0.9rem 1rem;
        margin-top: 0.5rem;
    }
    .stats-label {
        font-family: var(--font-mono);
        font-size: 0.62rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 1.5px;
        margin-bottom: 0.6rem;
    }
    .stats-file {
        font-family: var(--font-mono);
        font-size: 0.72rem;
        color: var(--accent);
        background: var(--accent-glow2);
        border: 1px solid var(--border-bright);
        border-radius: var(--radius-sm);
        padding: 3px 8px;
        display: inline-block;
        margin-bottom: 4px;
        max-width: 100%;
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
    }
    .stats-count {
        font-size: 0.75rem;
        color: var(--text-secondary);
        margin-top: 6px;
    }
    .stats-count span {
        color: var(--success);
        font-weight: 600;
    }

    /* Guardrails info */
    .guardrails {
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-radius: var(--radius-md);
        padding: 0.9rem 1rem;
        margin-top: 0.8rem;
    }
    .guardrails-title {
        font-family: var(--font-mono);
        font-size: 0.62rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 1.5px;
        margin-bottom: 0.6rem;
    }
    .guardrail-item {
        font-size: 0.72rem;
        color: var(--text-secondary);
        padding: 3px 0;
        display: flex;
        align-items: center;
        gap: 6px;
    }
    .guardrail-dot {
        width: 5px;
        height: 5px;
        border-radius: 50%;
        background: var(--accent);
        flex-shrink: 0;
    }

    /* ── Header ── */
    .page-header {
        display: flex;
        align-items: flex-start;
        gap: 16px;
        margin-bottom: 2rem;
        padding-bottom: 1.5rem;
        border-bottom: 1px solid var(--border);
    }
    .header-icon-wrap {
        width: 52px;
        height: 52px;
        border-radius: 12px;
        background: var(--accent-glow);
        border: 1px solid rgba(0,212,255,0.2);
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.6rem;
        flex-shrink: 0;
    }
    .header-title {
        font-family: var(--font-mono);
        font-size: 1.75rem;
        font-weight: 600;
        color: var(--text-primary);
        letter-spacing: -0.5px;
        line-height: 1;
        margin-bottom: 6px;
    }
    .header-title span { color: var(--accent); }
    .header-sub {
        font-size: 0.82rem;
        color: var(--text-secondary);
        font-weight: 400;
    }
    .header-pills {
        display: flex;
        gap: 8px;
        margin-top: 8px;
        flex-wrap: wrap;
    }
    .header-pill {
        font-family: var(--font-mono);
        font-size: 0.62rem;
        color: var(--text-muted);
        border: 1px solid var(--border);
        border-radius: 20px;
        padding: 2px 10px;
        letter-spacing: 0.5px;
    }

    /* ── Tabs ── */
    .stTabs [data-baseweb="tab-list"] {
        background: transparent !important;
        border-bottom: 1px solid var(--border) !important;
        gap: 0 !important;
        padding: 0 !important;
    }
    .stTabs [data-baseweb="tab"] {
        font-family: var(--font-mono) !important;
        font-size: 0.78rem !important;
        font-weight: 500 !important;
        color: var(--text-muted) !important;
        padding: 0.7rem 1.2rem !important;
        border-bottom: 2px solid transparent !important;
        background: transparent !important;
        transition: color 0.2s !important;
        letter-spacing: 0.3px;
    }
    .stTabs [aria-selected="true"] {
        color: var(--accent) !important;
        border-bottom-color: var(--accent) !important;
    }
    .stTabs [data-baseweb="tab-panel"] {
        padding: 1.5rem 0 0 !important;
    }

    /* ── Buttons ── */
    .stButton > button {
        background: var(--accent) !important;
        color: #000 !important;
        border: none !important;
        border-radius: var(--radius-sm) !important;
        font-family: var(--font-mono) !important;
        font-weight: 600 !important;
        font-size: 0.8rem !important;
        padding: 0.55rem 1.4rem !important;
        letter-spacing: 0.3px;
        transition: all 0.15s ease !important;
        width: 100% !important;
    }
    .stButton > button:hover {
        background: #33ddff !important;
        transform: translateY(-1px);
        box-shadow: 0 4px 20px var(--accent-glow) !important;
    }
    .stButton > button:active {
        transform: translateY(0) !important;
    }

    /* ── Text Input ── */
    .stTextInput > div > div > input {
        background: var(--bg-card) !important;
        color: var(--text-primary) !important;
        border: 1.5px solid var(--border-bright) !important;
        border-radius: var(--radius-md) !important;
        font-family: var(--font-main) !important;
        font-size: 0.9rem !important;
        padding: 0.7rem 1rem !important;
        transition: border-color 0.2s !important;
    }
    .stTextInput > div > div > input:focus {
        border-color: var(--accent) !important;
        box-shadow: 0 0 0 3px var(--accent-glow) !important;
    }
    .stTextInput > div > div > input::placeholder {
        color: var(--text-muted) !important;
    }

    /* ── Answer Card ── */
    .answer-card {
        background: var(--bg-card);
        border: 1px solid var(--border-bright);
        border-top: 3px solid var(--accent);
        border-radius: var(--radius-lg);
        padding: 1.4rem 1.6rem;
        margin: 1.2rem 0;
        box-shadow: 0 8px 32px rgba(0,0,0,0.3);
        animation: slideIn 0.3s ease;
    }
    @keyframes slideIn {
        from { opacity: 0; transform: translateY(8px); }
        to   { opacity: 1; transform: translateY(0); }
    }
    .answer-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 0.9rem;
    }
    .answer-label {
        font-family: var(--font-mono);
        font-size: 0.62rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 1.5px;
    }
    .answer-text {
        color: var(--text-primary);
        font-size: 0.94rem;
        line-height: 1.8;
        font-weight: 400;
    }

    /* ── Confidence Badges ── */
    .badge {
        display: inline-flex;
        align-items: center;
        gap: 5px;
        font-family: var(--font-mono);
        font-size: 0.7rem;
        font-weight: 600;
        padding: 3px 10px;
        border-radius: 20px;
        letter-spacing: 0.3px;
    }
    .badge-dot { width: 6px; height: 6px; border-radius: 50%; }
    .badge-high {
        background: var(--success-bg);
        color: var(--success);
        border: 1px solid rgba(0, 229, 160, 0.3);
    }
    .badge-high .badge-dot { background: var(--success); }
    .badge-mid {
        background: var(--warn-bg);
        color: var(--warn);
        border: 1px solid rgba(255, 179, 71, 0.3);
    }
    .badge-mid .badge-dot { background: var(--warn); }
    .badge-low {
        background: var(--danger-bg);
        color: var(--danger);
        border: 1px solid rgba(255, 92, 92, 0.3);
    }
    .badge-low .badge-dot { background: var(--danger); }
    .badge-na {
        background: rgba(122,143,168,0.08);
        color: var(--text-secondary);
        border: 1px solid var(--border);
    }

    /* ── Source Chunks ── */
    .source-chunk {
        background: var(--bg-elevated);
        border: 1px solid var(--border);
        border-left: 3px solid var(--border-bright);
        border-radius: var(--radius-md);
        padding: 0.9rem 1.1rem;
        margin: 0.5rem 0;
        font-family: var(--font-mono);
        font-size: 0.74rem;
        color: var(--text-secondary);
        line-height: 1.7;
        transition: border-left-color 0.2s;
    }
    .source-chunk:hover { border-left-color: var(--accent-dim); }
    .source-chunk-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-bottom: 8px;
    }
    .source-chunk-label {
        font-size: 0.62rem;
        color: var(--accent);
        text-transform: uppercase;
        letter-spacing: 1px;
    }
    .source-chunk-score {
        font-size: 0.62rem;
        color: var(--text-muted);
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: 4px;
        padding: 1px 7px;
    }

    /* ── Refusal Card ── */
    .refusal-card {
        background: var(--danger-bg);
        border: 1px solid rgba(255,92,92,0.2);
        border-left: 3px solid var(--danger);
        border-radius: var(--radius-md);
        padding: 1.1rem 1.4rem;
        color: #ffb3b3;
        font-size: 0.88rem;
        line-height: 1.7;
        animation: slideIn 0.3s ease;
    }
    .refusal-card strong { color: var(--danger); }
    .refusal-icon { font-size: 1rem; margin-right: 6px; }

    /* ── Not Found Card ── */
    .notfound-card {
        background: var(--warn-bg);
        border: 1px solid rgba(255,179,71,0.2);
        border-left: 3px solid var(--warn);
        border-radius: var(--radius-md);
        padding: 1.1rem 1.4rem;
        color: #ffe0a8;
        font-size: 0.88rem;
        line-height: 1.7;
        animation: slideIn 0.3s ease;
    }
    .notfound-card strong { color: var(--warn); }

    /* ── Example Chips ── */
    .chips-section {
        margin-top: 1.5rem;
    }
    .chips-label {
        font-family: var(--font-mono);
        font-size: 0.62rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 1.5px;
        margin-bottom: 0.7rem;
    }

    /* Override chip buttons to look like chips */
    div[data-chip="true"] .stButton > button {
        background: var(--bg-elevated) !important;
        color: var(--text-secondary) !important;
        border: 1px solid var(--border-bright) !important;
        font-family: var(--font-mono) !important;
        font-size: 0.72rem !important;
        font-weight: 400 !important;
        padding: 0.35rem 0.9rem !important;
        border-radius: 20px !important;
        width: auto !important;
        white-space: nowrap;
        transition: all 0.15s !important;
    }
    div[data-chip="true"] .stButton > button:hover {
        background: var(--bg-card) !important;
        border-color: var(--accent-dim) !important;
        color: var(--accent) !important;
        transform: none !important;
        box-shadow: none !important;
    }

    /* ── History Items ── */
    .history-item {
        background: var(--bg-card);
        border: 1px solid var(--border);
        border-radius: var(--radius-md);
        padding: 1rem 1.2rem;
        margin-bottom: 0.6rem;
        transition: border-color 0.2s;
    }
    .history-item:hover { border-color: var(--border-bright); }
    .history-q {
        font-size: 0.85rem;
        font-weight: 600;
        color: var(--accent);
        margin-bottom: 5px;
        display: flex;
        align-items: flex-start;
        gap: 8px;
    }
    .history-q-icon { color: var(--text-muted); flex-shrink: 0; }
    .history-a {
        font-size: 0.8rem;
        color: var(--text-secondary);
        line-height: 1.6;
        padding-left: 18px;
    }
    .history-meta {
        display: flex;
        align-items: center;
        justify-content: flex-end;
        margin-top: 8px;
    }

    /* ── Divider ── */
    hr {
        border: none !important;
        border-top: 1px solid var(--border) !important;
        margin: 1.5rem 0 !important;
    }

    /* ── Section Heading ── */
    .section-heading {
        font-family: var(--font-mono);
        font-size: 0.68rem;
        color: var(--text-muted);
        text-transform: uppercase;
        letter-spacing: 2px;
        margin-bottom: 0.9rem;
        display: flex;
        align-items: center;
        gap: 8px;
    }
    .section-heading::after {
        content: '';
        flex: 1;
        height: 1px;
        background: var(--border);
    }

    /* ── Empty State ── */
    .empty-state {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 4rem 2rem;
        text-align: center;
    }
    .empty-icon {
        font-size: 3rem;
        margin-bottom: 1rem;
        opacity: 0.4;
    }
    .empty-title {
        font-family: var(--font-mono);
        font-size: 0.9rem;
        color: var(--text-secondary);
        margin-bottom: 0.4rem;
    }
    .empty-sub {
        font-size: 0.78rem;
        color: var(--text-muted);
    }

    /* ── JSON Display ── */
    div[data-testid="stJson"] {
        background: var(--bg-elevated) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-md) !important;
    }

    /* ── Spinner ── */
    .stSpinner > div { border-top-color: var(--accent) !important; }

    /* ── Expander ── */
    details {
        background: var(--bg-elevated) !important;
        border: 1px solid var(--border) !important;
        border-radius: var(--radius-md) !important;
        padding: 0 !important;
    }
    summary {
        font-family: var(--font-mono) !important;
        font-size: 0.78rem !important;
        color: var(--text-secondary) !important;
        padding: 0.7rem 1rem !important;
    }
    details[open] summary { border-bottom: 1px solid var(--border); }

    /* ── Process button — special override ── */
    .process-btn .stButton > button {
        background: transparent !important;
        color: var(--accent) !important;
        border: 1.5px solid var(--accent) !important;
        font-size: 0.78rem !important;
    }
    .process-btn .stButton > button:hover {
        background: var(--accent-glow) !important;
        box-shadow: 0 4px 16px var(--accent-glow) !important;
    }

    /* Field summary grid */
    .field-row {
        display: flex;
        align-items: center;
        gap: 10px;
        padding: 0.55rem 0;
        border-bottom: 1px solid var(--border);
        font-size: 0.82rem;
    }
    .field-row:last-child { border-bottom: none; }
    .field-name {
        color: var(--text-secondary);
        min-width: 150px;
        font-size: 0.78rem;
    }
    .field-value {
        font-family: var(--font-mono);
        font-size: 0.78rem;
        color: var(--text-primary);
    }
    .field-null {
        color: var(--text-muted);
        font-style: italic;
    }
    .field-check { color: var(--success); font-size: 0.75rem; }
    .field-empty { color: var(--text-muted); font-size: 0.75rem; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# TEXT EXTRACTION
# ─────────────────────────────────────────────
def extract_text_from_files(uploaded_files):
    text = ""
    for file in uploaded_files:
        file_type = file.name.split(".")[-1].lower()
        if file_type == "pdf":
            text += extract_pdf_text(file)
        elif file_type == "docx":
            text += extract_docx_text(file)
        elif file_type == "txt":
            text += extract_txt_text(file)
        else:
            st.warning(f"Unsupported file type: {file.name}")
    return text

def extract_docx_text(file):
    text = ""
    doc = DocxDocument(file)
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"
    for table in doc.tables:
        text += "\n"
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text += " | ".join(row_text) + "\n"
    return text

def extract_txt_text(file):
    return file.read().decode("utf-8").strip()

def extract_pdf_text(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(layout=True)
            if page_text:
                text += page_text + "\n"
    return text


# ─────────────────────────────────────────────
# CHUNKING & VECTOR STORE
# ─────────────────────────────────────────────
def get_text_chunks(text):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=120,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    return splitter.split_text(text)

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    return vector_store


# ─────────────────────────────────────────────
# CONFIDENCE PARSING
# ─────────────────────────────────────────────
def parse_confidence(text: str) -> int:
    patterns = [
        r"confidence[^\d]{0,15}(\d{1,3})",
        r"(\d{1,3})\s*(?:%|percent)\s*confident",
        r"\[(\d{1,3})\]",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = int(m.group(1))
            if 0 <= val <= 100:
                return val
    return None

def split_answer_and_confidence(raw: str):
    confidence = parse_confidence(raw)
    clean = re.sub(
        r"(confidence\s*(?:score)?[:\s\-]*\d{1,3}\s*%?\.?)",
        "",
        raw,
        flags=re.IGNORECASE
    ).strip()
    return clean, confidence

def confidence_badge_html(score: int) -> str:
    if score is None:
        return '<span class="badge badge-na"><span class="badge-dot"></span>N/A</span>'
    elif score >= 70:
        return f'<span class="badge badge-high"><span class="badge-dot"></span>{score}% confidence</span>'
    elif score >= 40:
        return f'<span class="badge badge-mid"><span class="badge-dot"></span>{score}% confidence</span>'
    else:
        return f'<span class="badge badge-low"><span class="badge-dot"></span>{score}% confidence</span>'


# ─────────────────────────────────────────────
# RETRIEVAL WITH SIMILARITY GUARD
# ─────────────────────────────────────────────
def retrieve_with_scores(question: str, vector_store):
    results = vector_store.similarity_search_with_score(question, k=4)
    docs = [r[0] for r in results]
    raw_scores = [r[1] for r in results]
    best_score = min(raw_scores)
    passed = best_score <= SIMILARITY_THRESHOLD
    display_scores = [max(0, round((1 - s / 2) * 100)) for s in raw_scores]  # heuristic: L2=0→100%, L2=2→0%
    return docs, display_scores, passed, best_score


# ─────────────────────────────────────────────
# LLM CHAIN
# ─────────────────────────────────────────────
def get_conversational_chain():
    prompt_template = """
You are a logistics document assistant. Answer ONLY from the provided context.

Rules:
- If the answer is clearly present: answer precisely and return a confidence score (0-100).
- If the answer is partially present: answer what you can, note what's missing, lower confidence.
- If the answer is NOT in the context at all: respond exactly with "Not found in document."
- Do NOT generalize or use outside knowledge.
- Always end your response with exactly this line:
  Confidence: <number between 0 and 100>

Context:
{context}

Question:
{question}

Answer:
"""
    model = ChatGoogleGenerativeAI(model="models/gemma-3-27b-it", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain


# ─────────────────────────────────────────────
# Q&A HANDLER
# ─────────────────────────────────────────────
def handle_question(user_question: str, vector_store):
    docs, display_scores, passed_threshold, best_l2 = retrieve_with_scores(
        user_question, vector_store
    )

    if not passed_threshold:
        retrieval_score = display_scores[0] if display_scores else 0
        st.markdown(f"""
        <div class="refusal-card">
            <span class="refusal-icon">⚠</span>
            <strong>Low retrieval confidence — answer refused</strong><br>
            Best matching chunk has L2 distance <strong>{best_l2:.3f}</strong> (threshold: {SIMILARITY_THRESHOLD}).
            Displayed similarity: <strong>{retrieval_score}%</strong>.
            This question likely refers to content not present in the uploaded document.
        </div>
        """, unsafe_allow_html=True)
        _record_history(user_question, "⚠️ Refused — low retrieval similarity.", None)
        return

    chain = get_conversational_chain()
    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )
    raw_answer = response["output_text"]

    if "not found in document" in raw_answer.lower():
        st.markdown("""
        <div class="notfound-card">
            <span class="refusal-icon">🔍</span>
            <strong>Not found in document</strong><br>
            The uploaded document does not contain information to answer this question.
        </div>
        """, unsafe_allow_html=True)
        _record_history(user_question, "Not found in document.", 0)
        return

    clean_answer, confidence = split_answer_and_confidence(raw_answer)

    if confidence is not None and confidence < LOW_CONFIDENCE_CUTOFF:
        st.markdown(f"""
        <div class="refusal-card">
            <span class="refusal-icon">⚠</span>
            <strong>Answer confidence too low ({confidence}%) — withheld</strong><br>
            The model was not confident enough based on document content.
            Try rephrasing your question or verify this information exists in the document.
        </div>
        """, unsafe_allow_html=True)
        _record_history(user_question, f"Withheld — confidence {confidence}%", confidence)
        return

    badge = confidence_badge_html(confidence)
    st.markdown(f"""
    <div class="answer-card">
        <div class="answer-header">
            <span class="answer-label">Response</span>
            {badge}
        </div>
        <div class="answer-text">{clean_answer}</div>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("View source chunks"):
        for i, (doc, score) in enumerate(zip(docs, display_scores)):
            st.markdown(f"""
            <div class="source-chunk">
                <div class="source-chunk-header">
                    <span class="source-chunk-label">Chunk {i+1}</span>
                    <span class="source-chunk-score">{score}% match</span>
                </div>
                {doc.page_content}
            </div>
            """, unsafe_allow_html=True)

    _record_history(user_question, clean_answer, confidence)


def _record_history(question, answer, confidence):
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    st.session_state.chat_history.insert(0, {
        "user": question,
        "bot": answer,
        "confidence": confidence
    })


# ─────────────────────────────────────────────
# STRUCTURED EXTRACTION
# ─────────────────────────────────────────────
EXTRACTION_QUESTION = """
Extract the following fields strictly from the document. Return ONLY valid JSON, no explanation, no markdown code fences.
If a field is not found return null for that field.

{
  "shipment_id": null,
  "shipper": null,
  "consignee": null,
  "pickup_datetime": null,
  "delivery_datetime": null,
  "equipment_type": null,
  "mode": null,
  "rate": null,
  "currency": null,
  "weight": null,
  "carrier_name": null
}
"""

def run_extraction(vector_store):
    docs, display_scores, passed_threshold, _ = retrieve_with_scores(
        EXTRACTION_QUESTION, vector_store
    )
    all_docs, _, _, _ = retrieve_with_scores("shipment shipper consignee rate carrier", vector_store)
    docs = list({d.page_content: d for d in docs + all_docs}.values())

    chain = get_conversational_chain()
    response = chain(
        {"input_documents": docs, "question": EXTRACTION_QUESTION},
        return_only_outputs=True
    )

    raw = response["output_text"].strip()
    raw = re.sub(r"^```(?:json)?", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"```$", "", raw).strip()

    json_match = re.search(r"\{.*\}", raw, re.DOTALL)
    if json_match:
        raw = json_match.group(0)

    try:
        parsed = json.loads(raw)
        return parsed, True
    except json.JSONDecodeError:
        return raw, False


# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div class="sidebar-brand">
            <span class="sidebar-brand-icon">📦</span>
            <div>
                <div class="sidebar-brand-text">TMS Intelligence</div>
                <div class="sidebar-brand-sub">LOGISTICS · DOCUMENT · AI</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown('<div style="height:4px"></div>', unsafe_allow_html=True)

        uploaded_files = st.file_uploader(
            "Upload Document(s)",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            help="Supported: PDF, DOCX, TXT · Max 200MB per file"
        )

        st.markdown('<div style="height:8px"></div>', unsafe_allow_html=True)

        st.markdown('<div class="process-btn">', unsafe_allow_html=True)
        if st.button("⚙  Process Documents"):
            if uploaded_files:
                with st.spinner("Building index..."):
                    raw_text = extract_text_from_files(uploaded_files)
                    if not raw_text.strip():
                        st.error("No text could be extracted.")
                    else:
                        chunks = get_text_chunks(raw_text)
                        st.session_state.vector_store = get_vector_store(chunks)
                        st.session_state.raw_text = raw_text
                        st.session_state.chunk_count = len(chunks)
                        st.session_state.file_names = [f.name for f in uploaded_files]
                        st.session_state.chat_history = []
                        st.success(f"✓ Indexed {len(chunks)} chunks from {len(uploaded_files)} file(s)")
            else:
                st.warning("Please upload at least one file.")
        st.markdown('</div>', unsafe_allow_html=True)

        if st.session_state.get("vector_store"):
            names = st.session_state.get("file_names", [])
            chunks = st.session_state.get("chunk_count", 0)
            files_html = "".join(f'<div class="stats-file">{n}</div>' for n in names)
            st.markdown(f"""
            <div class="stats-box">
                <div class="stats-label">Index Active</div>
                {files_html}
                <div class="stats-count"><span>{chunks}</span> chunks indexed</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("""
        <div class="guardrails">
            <div class="guardrails-title">Active Guardrails</div>
            <div class="guardrail-item"><span class="guardrail-dot"></span>Retrieval similarity threshold</div>
            <div class="guardrail-item"><span class="guardrail-dot"></span>LLM confidence scoring</div>
            <div class="guardrail-item"><span class="guardrail-dot"></span>"Not found" passthrough</div>
            <div class="guardrail-item"><span class="guardrail-dot"></span>Low-confidence refusal</div>
        </div>
        """, unsafe_allow_html=True)

    return uploaded_files


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    if 'vector_store' not in st.session_state:
        st.session_state.vector_store = None
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    render_sidebar()

    # Header
    st.markdown("""
    <div class="page-header">
        <div class="header-icon-wrap">📦</div>
        <div>
            <div class="header-title">TMS Doc <span>Intelligence</span></div>
            <div class="header-sub">Upload logistics documents · Query with AI · Extract structured data</div>
            <div class="header-pills">
                <span class="header-pill">RAG · FAISS</span>
                <span class="header-pill">Gemma 3 27B</span>
                <span class="header-pill">Confidence Scoring</span>
                <span class="header-pill">PDF · DOCX · TXT</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.vector_store:
        st.markdown("""
        <div class="empty-state">
            <div class="empty-icon">⬆</div>
            <div class="empty-title">No document indexed</div>
            <div class="empty-sub">Upload a PDF, DOCX, or TXT file from the sidebar and click Process Documents</div>
        </div>
        """, unsafe_allow_html=True)
        return

    tab_qa, tab_extract, tab_history = st.tabs(["  💬  Ask Questions  ", "  🗂  Extract Data  ", "  📜  History  "])

    # ══ TAB 1 — Q&A ══
    with tab_qa:
        with st.form("qa_form", clear_on_submit=True):
            user_question = st.text_input(
                "Question",
                placeholder="e.g. What is the carrier rate? When is pickup scheduled?",
                label_visibility="collapsed"
            )
            submitted = st.form_submit_button("Ask →")

        if submitted and user_question.strip():
            with st.spinner("Retrieving context and generating answer..."):
                handle_question(user_question.strip(), st.session_state.vector_store)

        # Example chips
        st.markdown("""
        <div class="chips-section">
            <div class="chips-label">Example questions</div>
        </div>
        """, unsafe_allow_html=True)

        examples = [
            "What is the carrier rate?",
            "Who is the consignee?",
            "When is pickup scheduled?",
            "What equipment type is required?",
            "What is the shipper's name?",
            "What is the delivery date?",
        ]
        cols = st.columns(3)
        for i, ex in enumerate(examples):
            with cols[i % 3]:
                st.markdown('<div data-chip="true">', unsafe_allow_html=True)
                if st.button(ex, key=f"ex_{i}"):
                    with st.spinner("Retrieving context and generating answer..."):
                        handle_question(ex, st.session_state.vector_store)
                st.markdown('</div>', unsafe_allow_html=True)

    # ══ TAB 2 — EXTRACTION ══
    with tab_extract:
        st.markdown("""
        <div class="section-heading">Structured Extraction</div>
        """, unsafe_allow_html=True)
        st.markdown('<p style="color:var(--text-secondary); font-size:0.85rem; margin-bottom:1rem;">Automatically extract all standard TMS fields. Missing fields return as <code style="font-family:var(--font-mono); font-size:0.78rem; color:var(--text-muted);">null</code>.</p>', unsafe_allow_html=True)

        if st.button("🔍  Run Extraction"):
            with st.spinner("Extracting structured fields..."):
                result, success = run_extraction(st.session_state.vector_store)

            if success:
                st.success("Extraction complete")

                col_json, col_fields = st.columns([1, 1])
                with col_json:
                    st.markdown('<div class="section-heading">Raw JSON</div>', unsafe_allow_html=True)
                    st.json(result)

                with col_fields:
                    st.markdown('<div class="section-heading">Field Summary</div>', unsafe_allow_html=True)
                    fields = [
                        ("Shipment ID", result.get("shipment_id")),
                        ("Shipper", result.get("shipper")),
                        ("Consignee", result.get("consignee")),
                        ("Pickup", result.get("pickup_datetime")),
                        ("Delivery", result.get("delivery_datetime")),
                        ("Equipment", result.get("equipment_type")),
                        ("Mode", result.get("mode")),
                        ("Rate", result.get("rate")),
                        ("Currency", result.get("currency")),
                        ("Weight", result.get("weight")),
                        ("Carrier", result.get("carrier_name")),
                    ]
                    rows_html = ""
                    for field, value in fields:
                        if value:
                            rows_html += f'<div class="field-row"><span class="field-check">✓</span><span class="field-name">{field}</span><span class="field-value">{value}</span></div>'
                        else:
                            rows_html += f'<div class="field-row"><span class="field-empty">○</span><span class="field-name">{field}</span><span class="field-value field-null">null</span></div>'
                    st.markdown(f'<div style="background:var(--bg-card); border:1px solid var(--border); border-radius:10px; padding:0.8rem 1rem;">{rows_html}</div>', unsafe_allow_html=True)
            else:
                st.warning("Could not parse JSON from model response.")
                st.code(result, language="text")

    # ══ TAB 3 — HISTORY ══
    with tab_history:
        st.markdown('<div class="section-heading">Conversation History</div>', unsafe_allow_html=True)

        if not st.session_state.chat_history:
            st.markdown("""
            <div class="empty-state">
                <div class="empty-icon" style="font-size:2rem;">📜</div>
                <div class="empty-title">No history yet</div>
                <div class="empty-sub">Ask questions in the Q&A tab to populate history</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            if st.button("🗑  Clear History"):
                st.session_state.chat_history = []
                st.rerun()

            for item in st.session_state.chat_history:
                conf = item.get("confidence")
                badge = confidence_badge_html(conf) if conf is not None else ""
                preview = item['bot'][:280] + ('…' if len(item['bot']) > 280 else '')
                st.markdown(f"""
                <div class="history-item">
                    <div class="history-q"><span class="history-q-icon">›</span>{item['user']}</div>
                    <div class="history-a">{preview}</div>
                    <div class="history-meta">{badge}</div>
                </div>
                """, unsafe_allow_html=True)


if __name__ == "__main__":

    main()


